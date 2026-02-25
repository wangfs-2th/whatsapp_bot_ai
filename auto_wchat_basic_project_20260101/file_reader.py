"""
文件读取模块 - 支持读取doc和excel文件
"""
import os
import logging
from typing import Optional, Dict
from docx import Document
import pandas as pd

logger = logging.getLogger(__name__)


class FileReader:
    """文件读取器"""
    
    def __init__(self):
        self.supported_formats = ['.docx', '.doc', '.xlsx', '.xls']
    
    def read_file(self, file_path: str) -> Optional[str]:
        """读取文件内容"""
        # 转换为绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.abspath(file_path)
        
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return None
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.docx':
                return self._read_docx(file_path)
            elif file_ext == '.doc':
                return self._read_doc(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._read_excel(file_path)
            else:
                logger.error(f"不支持的文件格式: {file_ext}")
                return None
        except Exception as e:
            logger.error(f"读取文件失败: {e}")
            return None
    
    def _read_docx(self, file_path: str) -> str:
        """读取Word文档 (.docx格式)"""
        doc = Document(file_path)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    content.append(row_text)
        
        return '\n'.join(content)
    
    def _read_doc(self, file_path: str) -> str:
        """读取旧版Word文档 (.doc格式)"""
        # python-docx 不支持 .doc 格式，需要其他方法
        # 方法1: 尝试使用 textract (如果可用)
        try:
            import textract
            content = textract.process(file_path).decode('utf-8')
            return content
        except ImportError:
            logger.debug("textract 库未安装，尝试其他方法")
        except Exception as e:
            logger.debug(f"使用 textract 读取失败: {e}")
        
        # 方法2: 尝试使用 win32com (Windows only)
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path)
            content = doc.Content.Text
            doc.Close()
            word.Quit()
            return content
        except ImportError:
            logger.debug("win32com 库未安装，尝试其他方法")
        except Exception as e:
            logger.debug(f"使用 win32com 读取失败: {e}")
        
        # 方法3: 提示用户转换为 .docx
        error_msg = (
            "无法读取 .doc 格式文件。python-docx 库只支持 .docx 格式。"
            "解决方案：\n"
            "1. 将文件转换为 .docx 格式后重新上传（推荐）\n"
            "2. 安装 textract 库以支持 .doc 格式：pip install textract\n"
            "3. 在 Windows 系统上安装 pywin32：pip install pywin32"
        )
        logger.error(f"读取 .doc 文件失败: {file_path}")
        raise Exception(error_msg)
    
    def _read_excel(self, file_path: str) -> str:
        """读取Excel文件"""
        try:
            df = pd.read_excel(file_path, sheet_name=None)
            
            content = []
            for sheet_name, sheet_df in df.items():
                content.append(f"\n=== 工作表: {sheet_name} ===\n")
                
                for index, row in sheet_df.iterrows():
                    row_text = ' | '.join([str(val) if pd.notna(val) else '' for val in row.values])
                    if row_text.strip():
                        content.append(row_text)
            
            return '\n'.join(content)
        
        except Exception as e:
            logger.warning(f"使用pandas读取失败: {e}")
            return ""
    
    def extract_key_info(self, content: str, max_length: int = 2000) -> str:
        """提取关键信息并整理"""
        if not content:
            return ""
        
        lines = [line.strip() for line in content.split('\n') if line.strip()]
        cleaned_content = '\n'.join(lines)
        
        if len(cleaned_content) > max_length:
            cleaned_content = cleaned_content[:max_length] + "..."
        
        return cleaned_content
    
    def process_file_for_ai(self, file_path: str, max_length: int = 2000) -> Optional[str]:
        """读取文件并整理内容，用于AI回复"""
        content = self.read_file(file_path)
        if not content:
            return None
        
        return self.extract_key_info(content, max_length)


class ContentManager:
    """内容管理器"""
    
    def __init__(self):
        self.file_reader = FileReader()
        self.loaded_contents: Dict[str, str] = {}
    
    def load_file(self, file_path: str) -> bool:
        """加载文件内容"""
        content = self.file_reader.process_file_for_ai(file_path)
        if content:
            self.loaded_contents[file_path] = content
            logger.info(f"文件已加载: {file_path}")
            return True
        return False
    
    def get_content(self, file_path: str) -> Optional[str]:
        """获取文件内容"""
        return self.loaded_contents.get(file_path)
    
    def get_all_content(self) -> str:
        """获取所有已加载的内容"""
        return '\n\n'.join(self.loaded_contents.values())
    
    def remove_file(self, file_path: str):
        """移除文件"""
        if file_path in self.loaded_contents:
            del self.loaded_contents[file_path]
            logger.info(f"文件已移除: {file_path}")
    
    def clear_all(self):
        """清空所有内容"""
        self.loaded_contents.clear()
        logger.info("所有内容已清空")

